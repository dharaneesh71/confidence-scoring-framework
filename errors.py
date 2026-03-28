import subprocess
subprocess.run(["pip", "install", "python-docx", "-q"])

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ─────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   kwargs.get('val', 'single'))
        b.set(qn('w:sz'),    kwargs.get('sz', '4'))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), kwargs.get('color', 'DDDDDD'))
        tcBorders.append(b)
    tcPr.append(tcBorders)

# ── Colour palette ─────────────────────────────────────────────
RED_BG    = 'FFF0F0'
RED_BADGE = 'D32F2F'
ORA_BG    = 'FFF8F0'
ORA_BADGE = 'E65100'
YEL_BG    = 'FFFFF0'
YEL_BADGE = 'F57F17'
HDR_BG    = '1E3A5F'
ALT_BG    = 'F4F7FB'
CODE_BG   = 'F1F3F5'

# ══════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(title_para, '1E3A5F')
run = title_para.add_run('CONFID.AI — Code Review Bug Report')
run.bold      = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub_para, '1E3A5F')
run2 = sub_para.add_run('Repository: dharaneesh71/confidence-scoring-framework  |  Reviewed: March 28, 2026')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0xB0, 0xC8, 0xE8)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ══════════════════════════════════════════════════════════════
h = doc.add_heading('Summary', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

tbl = doc.add_table(rows=2, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

headers = ['Severity', '# Issues', 'Files Affected', 'Status']
data    = [
    ('🔴 Critical', '6',  'database.py, endpoints.py, llama_service.py, security.py', 'Fix Immediately'),
    ('🟠 Major',    '8',  'main.py, security.py, endpoints.py, chroma_service.py, AdminPage.js, conftest.py', 'Fix Before Release'),
    ('🟡 Minor',    '8',  'database.py, llama_service.py, config.py, App.js, api.js, Sidebar.js, tests/', 'Fix in Next Sprint'),
]

hdr_cells = tbl.rows[0].cells
for i, h_txt in enumerate(headers):
    hdr_cells[i].text = h_txt
    shade_cell(hdr_cells[i], HDR_BG)
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

# Delete the auto-created second row
tbl._tbl.remove(tbl.rows[1]._tr)

SUMMARY = [
    ('🔴  Critical', '6',  'database.py, endpoints.py,\nllama_service.py, security.py',  '⚠ Fix Immediately',  'FFF0F0', 'D32F2F'),
    ('🟠  Major',    '8',  'main.py, security.py, AdminPage.js,\nchroma_service.py, conftest.py', 'Fix Before Release', 'FFF8F0', 'E65100'),
    ('🟡  Minor',    '8',  'config.py, App.js, api.js,\nSidebar.js, tests/',             'Fix in Next Sprint',  'FFFFF0', 'F57F17'),
]
for sev, cnt, files, status, bg, _ in SUMMARY:
    row = tbl.add_row()
    vals = [sev, cnt, files, status]
    for i, val in enumerate(vals):
        row.cells[i].text = val
        shade_cell(row.cells[i], bg)
        p = row.cells[i].paragraphs[0]
        p.runs[0].font.size = Pt(9)
        if i == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# HELPER: add an issue block
# ══════════════════════════════════════════════════════════════
def add_issue(number, severity, filename, title, bug_desc, bug_code, fix_desc, fix_code, bg_color, badge_color_hex):
    # Badge line
    badge_para = doc.add_paragraph()
    shade_paragraph(badge_para, bg_color)
    badge_run = badge_para.add_run(f'  #{number}  {severity}   {filename}')
    badge_run.bold      = True
    badge_run.font.size = Pt(10)
    badge_run.font.color.rgb = RGBColor(*bytes.fromhex(badge_color_hex))

    # Title
    t_para = doc.add_paragraph()
    shade_paragraph(t_para, bg_color)
    t_run = t_para.add_run(f'     {title}')
    t_run.bold      = True
    t_run.font.size = Pt(11)

    # Bug description
    bd_para = doc.add_paragraph()
    shade_paragraph(bd_para, bg_color)
    bd_para.paragraph_format.left_indent = Inches(0.3)
    bd_run = bd_para.add_run(f'BUG:  {bug_desc}')
    bd_run.font.size = Pt(9.5)

    # Bug code (if any)
    if bug_code:
        bc_para = doc.add_paragraph()
        shade_paragraph(bc_para, CODE_BG)
        bc_para.paragraph_format.left_indent = Inches(0.4)
        bc_run = bc_para.add_run(bug_code)
        bc_run.font.name = 'Courier New'
        bc_run.font.size = Pt(8.5)
        bc_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Fix description
    fx_para = doc.add_paragraph()
    shade_paragraph(fx_para, bg_color)
    fx_para.paragraph_format.left_indent = Inches(0.3)
    fx_run = fx_para.add_run(f'FIX:  {fix_desc}')
    fx_run.font.size = Pt(9.5)
    fx_run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    fx_run.bold = True

    # Fix code (if any)
    if fix_code:
        fxc_para = doc.add_paragraph()
        shade_paragraph(fxc_para, CODE_BG)
        fxc_para.paragraph_format.left_indent = Inches(0.4)
        fxc_run = fxc_para.add_run(fix_code)
        fxc_run.font.name = 'Courier New'
        fxc_run.font.size = Pt(8.5)
        fxc_run.font.color.rgb = RGBColor(0x00, 0x60, 0x00)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# SECTION HEADER HELPER
# ══════════════════════════════════════════════════════════════
def section_header(text, hex_bg, hex_fg='FFFFFF'):
    p = doc.add_paragraph()
    shade_paragraph(p, hex_bg)
    r = p.add_run(f'  {text}')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(*bytes.fromhex(hex_fg))
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# SECTION 1 — CRITICAL
# ══════════════════════════════════════════════════════════════
section_header('🔴  SECTION 1 — CRITICAL BUGS  (App will break or data will be lost)', 'B71C1C')

add_issue(1, '🔴 CRITICAL', 'backend/core/database.py',
    'ChatHistory model missing explanation & citations columns',
    'The ChatHistory ORM model has no explanation or citations columns. Every query response generates these values but they are immediately discarded — citations are permanently lost after the response is sent.',
    'class ChatHistory(Base):\n    question         = Column(String)\n    answer           = Column(String)\n    confidence_score = Column(Float)\n    # ← explanation column MISSING\n    # ← citations column MISSING',
    'Add both columns with Text type. Then DELETE confid_ai.db and restart — SQLite cannot auto-add columns to existing tables.',
    'from sqlalchemy import Text\n\nexplanation = Column(Text, nullable=True)\ncitations   = Column(Text, nullable=True)',
    RED_BG, 'B71C1C')

add_issue(2, '🔴 CRITICAL', 'backend/api/routes/endpoints.py  →  submit_query()',
    'explanation and citations never saved when inserting ChatHistory',
    'Even after fix #1, the ChatHistory constructor call in submit_query() still omits explanation and citations. They are computed but silently thrown away.',
    'history_entry = ChatHistory(\n    user_id=current_user.id,\n    session_id=session_id,\n    question=request.question,\n    answer=answer,\n    confidence_score=confidence_score\n    # ← explanation=explanation   MISSING\n    # ← citations=json.dumps(citations)  MISSING\n)',
    'Add both fields to the constructor. Also add "import json" at the top of endpoints.py.',
    'import json\n\nhistory_entry = ChatHistory(\n    user_id=current_user.id,\n    session_id=session_id,\n    question=request.question,\n    answer=answer,\n    confidence_score=confidence_score,\n    explanation=explanation,\n    citations=json.dumps(citations)   # serialize list → JSON string\n)',
    RED_BG, 'B71C1C')

add_issue(3, '🔴 CRITICAL', 'backend/api/routes/endpoints.py  →  get_session_details()',
    'Citations always deserialized as empty list on session reload',
    'Uses getattr(msg, "citations", []) which silently returns [] because the column never existed. Even after fix #1 the value stored is a JSON string — it must be parsed back.',
    '"citations": getattr(msg, "citations", []) or [],\n# getattr() returns [] forever → frontend never sees references on reload',
    'After fixing #1 and #2, deserialize the stored JSON string. Create a _serialize_message() helper to reuse in both get_session_details and get_session_analytics.',
    'def _serialize_message(msg):\n    return {\n        "history_id":       msg.id,\n        "question":         msg.question,\n        "answer":           msg.answer,\n        "confidence_score": msg.confidence_score,\n        "timestamp":        msg.timestamp,\n        "explanation":      msg.explanation,\n        "citations":        json.loads(msg.citations or "[]"),\n    }\n\n# Use in get_session_details:\nserialized_messages = [_serialize_message(m) for m in messages]',
    RED_BG, 'B71C1C')

add_issue(4, '🔴 CRITICAL', 'backend/services/llama_service.py  →  is_ready()',
    'is_ready() always returns True even when model failed to load',
    'The method unconditionally returns True. If LlamaService.__init__() crashes (e.g., GPU OOM, missing HuggingFace token), self.model is set to None but /api/status still reports "llm_ready": true — making it impossible to detect a broken service.',
    'def is_ready(self) -> bool:\n    return True   # ← always True regardless of self.model state',
    'Check the actual model and pipeline references.',
    'def is_ready(self) -> bool:\n    return self.model is not None and self.pipeline is not None',
    RED_BG, 'B71C1C')

add_issue(5, '🔴 CRITICAL', 'backend/api/routes/endpoints.py  →  get_session_analytics()',
    'Returns raw SQLAlchemy ORM objects — crashes with JSON serialization error',
    'The "messages" key in the return dict holds raw ORM objects. FastAPI cannot serialize these to JSON, causing a 500 error every time the analytics endpoint is called.',
    'return {\n    "session_id":         session_id,\n    "average_confidence": round(avg_score, 2),\n    "trend":              [{...}],\n    "messages":           messages   # ← raw ORM objects → runtime crash\n}',
    'Remove the "messages" key (unused by frontend) OR serialize using the _serialize_message() helper from fix #3.',
    'return {\n    "session_id":         session_id,\n    "average_confidence": round(avg_score, 2),\n    "total_interactions": len(messages),\n    "trend":              [{"turn": i+1, "score": round(s * 100, 1)}\n                           for i, s in enumerate(scores)],\n    # "messages" key removed — not needed by frontend\n}',
    RED_BG, 'B71C1C')

add_issue(6, '🔴 CRITICAL', 'backend/core/security.py',
    'SECRET_KEY is a hardcoded plaintext string committed to version control',
    'Any JWT token can be forged by anyone who reads this repository. There is no SECRET_KEY field in config.py so the setting can never be overridden via .env.',
    'SECRET_KEY = "YOUR_SUPER_SECRET_KEY_CHANGE_THIS_IN_PROD"\n# ← in git history forever, visible to anyone',
    'Add SECRET_KEY to config.py Settings and read it from environment in security.py.',
    '# config.py — add to Settings class:\nSECRET_KEY: str = ""  # Must be set in .env\n\n# security.py — replace hardcoded string:\nfrom core.config import settings\nSECRET_KEY = settings.SECRET_KEY\n\n# .env file:\nCONFID_SECRET_KEY=your-random-256-bit-secret-here',
    RED_BG, 'B71C1C')

# ══════════════════════════════════════════════════════════════
# SECTION 2 — MAJOR
# ══════════════════════════════════════════════════════════════
section_header('🟠  SECTION 2 — MAJOR BUGS  (Incorrect behaviour / runtime errors)', 'E65100')

add_issue(7, '🟠 MAJOR', 'backend/core/security.py',
    'Duplicate get_db() function — two DB session sources of truth',
    'get_db() is defined identically in both database.py and security.py. All endpoints import from core.security, which creates confusion and a maintenance risk if they ever diverge.',
    '# security.py — DUPLICATE:\ndef get_db():\n    db = SessionLocal()\n    try:\n        yield db\n    finally:\n        db.close()',
    'Delete get_db() from security.py entirely. It already exists in core/database.py and all files can import it from there.',
    '# security.py — remove the get_db function entirely\n# All endpoints should import from:\nfrom core.database import get_db',
    ORA_BG, 'E65100')

add_issue(8, '🟠 MAJOR', 'backend/services/llama_service.py',
    'Model name hardcoded — settings.LLAMA_MODEL_NAME is completely ignored',
    'The logger says it is loading settings.LLAMA_MODEL_NAME but the very next line overrides it with a hardcoded string. Changing LLAMA_MODEL_NAME in .env has zero effect.',
    'logger.info(f"Loading Llama model: {settings.LLAMA_MODEL_NAME}")  # logs "Meta-Llama-3.1-8B"\nmodel_name = "meta-llama/Llama-3.2-1B-Instruct"          # ← actually loads a different one',
    'Use the setting consistently.',
    'model_name = settings.LLAMA_MODEL_NAME',
    ORA_BG, 'E65100')

add_issue(9, '🟠 MAJOR', 'backend/main.py',
    'CORS allow_origins=["*"] ignores settings.ALLOWED_ORIGINS',
    'config.py defines a specific ALLOWED_ORIGINS list for security, but main.py bypasses it completely by passing ["*"] — allowing all origins in all environments.',
    'app.add_middleware(\n    CORSMiddleware,\n    allow_origins=["*"],    # ← ignores settings.ALLOWED_ORIGINS',
    'Use the configured list.',
    'app.add_middleware(\n    CORSMiddleware,\n    allow_origins=settings.ALLOWED_ORIGINS,',
    ORA_BG, 'E65100')

add_issue(10, '🟠 MAJOR', 'backend/core/security.py  →  get_current_active_admin()',
    'Admin privilege check raises HTTP 400 instead of 403 FORBIDDEN',
    'HTTP 400 means "Bad Request" (malformed input). A non-admin user making a valid request that they lack privileges for should get 403 FORBIDDEN, not 400.',
    'raise HTTPException(\n    status_code=400,    # ← wrong: 400 = bad request\n    detail="The user doesn\'t have enough privileges"\n)',
    'Use the correct HTTP status code.',
    'raise HTTPException(\n    status_code=status.HTTP_403_FORBIDDEN,\n    detail="The user doesn\'t have enough privileges"\n)',
    ORA_BG, 'E65100')

add_issue(11, '🟠 MAJOR', 'frontend/src/pages/AdminPage.js',
    'AdminPage container not scrollable — clipped by App.js overflow:hidden',
    'App.js sets overflow:"hidden" on the main content wrapper. AdminPage\'s <Container> has no compensating scroll. If the page content is taller than the viewport, everything below the fold is permanently inaccessible.',
    '// App.js content wrapper:\n<Box sx={{ overflow: "hidden", flex: 1, minHeight: 0 }}>\n\n// AdminPage — no scroll set:\n<Container maxWidth="lg" sx={{ my: 4 }}>   // ← gets clipped',
    'Add scrolling directly to the AdminPage container.',
    '// AdminPage.js:\n<Container maxWidth="lg"\n  sx={{ my: 4, height: "100%", overflowY: "auto" }}>',
    ORA_BG, 'E65100')

add_issue(12, '🟠 MAJOR', 'backend/api/routes/endpoints.py  (module level)',
    'All 4 services instantiated at module import time, not inside app lifespan',
    'pdf_processor, chroma_service, llama_service, and scoring_service are created when Python imports endpoints.py. This means the ~1GB Llama model downloads immediately on any import — including during tests, linting, or IDE indexing.',
    'router = APIRouter()\n\npdf_processor  = PDFProcessor()    # ← runs at import\nchroma_service = ChromaService()   # ← runs at import\nllama_service  = LlamaService()    # ← downloads 1GB at import\nscoring_service = ScoringService() # ← runs at import',
    'Move service init into main.py\'s lifespan() context manager and share via app.state.',
    '# main.py:\n@asynccontextmanager\nasync def lifespan(app: FastAPI):\n    app.state.llama    = LlamaService()\n    app.state.chroma   = ChromaService()\n    app.state.scoring  = ScoringService()\n    app.state.pdf      = PDFProcessor()\n    yield\n\n# endpoints.py: inject via Request\nfrom fastapi import Request\n@router.post("/query")\nasync def submit_query(req: Request, ...):\n    llama = req.app.state.llama',
    ORA_BG, 'E65100')

add_issue(13, '🟠 MAJOR', 'backend/services/chroma_service.py  →  search()',
    'math.exp(-distance) is a non-standard and misleading similarity conversion',
    'ChromaDB\'s default metric is L2 (Euclidean) distance. The exp(-distance) formula is an arbitrary mathematical choice that does not produce true similarity in [0,1] for L2 space, and can give inflated or inconsistent scores.',
    'distance = results["distances"][0][i]\nsimilarity = math.exp(-distance)   # ← arbitrary, non-standard',
    'Either switch the collection to cosine space and use 1-distance, or use 1/(1+distance) for L2 — both give a proper [0,1] similarity.',
    '# OPTION A — Cosine similarity (recommended with SentenceTransformers):\n# When creating the collection add:\n#   metadata={"hnsw:space": "cosine"}\n# Then convert:\nsimilarity = 1.0 - distance   # cosine distance [0,2] → similarity [0,1]\n\n# OPTION B — Keep L2, use proper formula:\nsimilarity = 1.0 / (1.0 + distance)',
    ORA_BG, 'E65100')

add_issue(14, '🟠 MAJOR', 'backend/conftest.py',
    'conftest.py is completely empty — tests load the real Llama model',
    'test_feedback.py imports from main, which imports endpoints.py, which instantiates LlamaService at module level (issue #12). Running pytest triggers a ~1GB model download with no mocking. Tests will timeout or crash on CI.',
    '# conftest.py — currently empty (0 bytes)',
    'Add mock fixtures to conftest.py that patch out heavy services before any test imports.',
    '# conftest.py:\nimport pytest\nfrom unittest.mock import MagicMock, patch\n\n@pytest.fixture(autouse=True)\ndef mock_services():\n    with patch("services.llama_service.LlamaService._initialize"),\\\n         patch("services.chroma_service.ChromaService._initialize"),\\\n         patch("services.scoring_service.ScoringService.__init__",\n               return_value=None):\n        yield',
    ORA_BG, 'E65100')

# ══════════════════════════════════════════════════════════════
# SECTION 3 — MINOR
# ══════════════════════════════════════════════════════════════
section_header('🟡  SECTION 3 — MINOR ISSUES  (Code quality, security hygiene, UX)', 'F57F17')

add_issue(15, '🟡 MINOR', 'backend/core/database.py',
    'Deprecated declarative_base import path',
    'SQLAlchemy 2.0 moved declarative_base to sqlalchemy.orm. The old import path from sqlalchemy.ext.declarative still works but emits deprecation warnings and will be removed.',
    'from sqlalchemy.ext.declarative import declarative_base  # deprecated',
    'Use the new import path.',
    'from sqlalchemy.orm import declarative_base',
    YEL_BG, 'F57F17')

add_issue(16, '🟡 MINOR', 'backend/services/llama_service.py  →  _initialize()',
    'Missing pad_token_id causes generation warnings and potential errors',
    'Llama tokenizers often lack a pad_token_id. Batched or padded generation calls will log warnings and may crash with older versions of transformers.',
    '# _initialize() — after pipeline is created, no pad_token_id is set',
    'Add the check immediately after the tokenizer is loaded.',
    '# Add directly after: self.tokenizer = AutoTokenizer.from_pretrained(...)\nif self.tokenizer.pad_token_id is None:\n    self.tokenizer.pad_token_id = self.tokenizer.eos_token_id',
    YEL_BG, 'F57F17')

add_issue(17, '🟡 MINOR', 'backend/core/config.py  +  createusers.py',
    'Hardcoded credentials committed to version control',
    'ADMIN_PASSWORD = "admin123" in config.py and plaintext passwords in createusers.py are checked into git. Even as defaults, they are security risks and will be in git history forever.',
    '# config.py:\nADMIN_PASSWORD: str = "admin123"\n\n# createusers.py:\ncreate_user("admin@example.com", "admin123", "ADMIN")\ncreate_user("user@example.com",  "user123",  "NORMAL USER")',
    'Remove defaults from code. Use .env exclusively. Add .env to .gitignore.',
    '# config.py — remove default value:\nADMIN_PASSWORD: str = ""\n\n# .env file (never commit this):\nADMIN_PASSWORD=your-secure-password-here\n\n# createusers.py — read from env or prompt:\nimport os\nadmin_pw = os.environ.get("ADMIN_PASSWORD") or input("Admin password: ")',
    YEL_BG, 'F57F17')

add_issue(18, '🟡 MINOR', 'frontend/src/App.js  +  QAPage.js  +  AdminPage.js  +  AuthContext.js',
    'Backend URL hardcoded as http://localhost:8000 in 8+ places',
    'Every fetch() call across the frontend hardcodes the localhost URL. This breaks the moment the app is deployed anywhere other than a local machine.',
    '# App.js:\nfetch("http://localhost:8000/api/history", ...)\n# AuthContext.js:\nfetch("http://localhost:8000/api/auth/login", ...)\n# And 6 more identical hardcoded URLs across QAPage.js and AdminPage.js',
    'Create a single shared constant using the REACT_APP_API_URL env var.',
    '# Create frontend/src/config.js:\nexport const API_BASE = process.env.REACT_APP_API_URL || "http://localhost:8000";\n\n# Then in every component:\nimport { API_BASE } from "../config";\nfetch(`${API_BASE}/api/history`, ...)\n\n# frontend/.env.local:\nREACT_APP_API_URL=http://localhost:8000',
    YEL_BG, 'F57F17')

add_issue(19, '🟡 MINOR', 'frontend/src/services/api.js',
    'Entire file is dead code — neither function is called, both have broken auth',
    'submitQuery() sends no Authorization header (would return 401). uploadDocument() uses Basic Auth instead of Bearer tokens. Neither function is imported or called anywhere — QAPage.js and AdminPage.js bypass this file entirely with direct fetch() calls.',
    '// submitQuery — no auth header:\nconst response = await api.post("/query", { question });\n// Would fail with 401 Unauthorized\n\n// uploadDocument — wrong auth type:\nauth: { username, password }  // Basic Auth, but API expects Bearer token',
    'Either wire up the correct Bearer token and use these functions throughout the app (replacing the direct fetch calls), or delete the file entirely.',
    '// Option A — Fix and use api.js (add auth interceptor):\napi.interceptors.request.use(config => {\n  const token = localStorage.getItem("token");\n  if (token) config.headers.Authorization = `Bearer ${token}`;\n  return config;\n});\n\n// Option B — Delete api.js and continue using fetch() directly.',
    YEL_BG, 'F57F17')

add_issue(20, '🟡 MINOR', 'frontend/src/components/Sidebar.js',
    'Sessions older than 7 days silently disappear from sidebar',
    'The groupSessions() function defines three groups: Today, Yesterday, and "Previous 7 Days". Any session with diffDays > 7 falls through the if/else chain and is never added to any group — it silently vanishes from the UI.',
    'const groups = { Today: [], Yesterday: [], "Previous 7 Days": [] };\n// ...\nif      (diffDays <= 0) groups.Today.push(s);\nelse if (diffDays === 1) groups.Yesterday.push(s);\nelse if (diffDays <= 7) groups["Previous 7 Days"].push(s);\n// ← sessions older than 7 days: dropped silently',
    'Add an "Older" catch-all group.',
    'const groups = { Today: [], Yesterday: [], "Previous 7 Days": [], "Older": [] };\n// ...\nif      (diffDays <= 0) groups.Today.push(s);\nelse if (diffDays === 1) groups.Yesterday.push(s);\nelse if (diffDays <= 7) groups["Previous 7 Days"].push(s);\nelse                    groups["Older"].push(s);  // ← catch-all',
    YEL_BG, 'F57F17')

add_issue(21, '🟡 MINOR', 'backend/tests/test_feedback.py',
    'Test ChatHistory fixture missing session_id — creates orphan DB record',
    'The test creates a ChatHistory row with no session_id. In production the session_id FK is always set. This means tests run against a different DB schema than production, potentially masking FK-related bugs.',
    'chat_entry = ChatHistory(\n    user_id=test_user.id,\n    question="Test Question",\n    answer="Test Answer",\n    confidence_score=0.9\n    # ← session_id missing\n)',
    'Create a parent ChatSession first in the fixture.',
    '@pytest.fixture(name="test_session")\ndef fixture_test_session(db_session, test_user):\n    session = ChatSession(user_id=test_user.id, title="Test Session")\n    db_session.add(session)\n    db_session.commit()\n    db_session.refresh(session)\n    return session\n\n# Then in test:\nchat_entry = ChatHistory(\n    user_id=test_user.id,\n    session_id=test_session.id,   # ← add this\n    question="Test Question",\n    answer="Test Answer",\n    confidence_score=0.9\n)',
    YEL_BG, 'F57F17')

add_issue(22, '🟡 MINOR', 'backend/services/scoring_service.py',
    'NLTK punkt and stopwords downloaded at startup but never used',
    'The scoring service downloads two NLTK datasets on every startup. Neither punkt tokenization nor stopwords removal is actually used anywhere in the file. This is a needless network call that will fail silently in air-gapped environments.',
    'import nltk\nnltk.download("punkt",     quiet=True)   # ← never used\nnltk.download("stopwords", quiet=True)   # ← never used',
    'Remove the NLTK import block entirely.',
    '# Delete these lines from scoring_service.py:\n# import nltk\n# try:\n#     nltk.download("punkt", quiet=True)\n#     nltk.download("stopwords", quiet=True)\n# except:\n#     pass',
    YEL_BG, 'F57F17')

# ══════════════════════════════════════════════════════════════
# RECOMMENDED FIX ORDER TABLE
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
h2 = doc.add_heading('Recommended Fix Order', level=1)
h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

order_tbl = doc.add_table(rows=1, cols=4)
order_tbl.style = 'Table Grid'
order_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, hdr in enumerate(['Order', 'Issue #', 'File', 'What to Fix']):
    cell = order_tbl.rows[0].cells[i]
    cell.text = hdr
    shade_cell(cell, HDR_BG)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

ORDER_ROWS = [
    ('1st', '#1', 'database.py',         'Add explanation & citations columns → delete DB → restart'),
    ('2nd', '#2', 'endpoints.py',        'Save explanation & citations in ChatHistory insert'),
    ('3rd', '#3', 'endpoints.py',        'Deserialize citations JSON in get_session_details'),
    ('4th', '#5', 'endpoints.py',        'Remove raw ORM objects from get_session_analytics return'),
    ('5th', '#6', 'security.py',         'Move SECRET_KEY to config.py + .env'),
    ('6th', '#4', 'llama_service.py',    'Fix is_ready() to check self.model is not None'),
    ('7th', '#7', 'security.py',         'Remove duplicate get_db() function'),
    ('8th', '#8', 'llama_service.py',    'Replace hardcoded model_name with settings.LLAMA_MODEL_NAME'),
    ('9th', '#9', 'main.py',             'Use settings.ALLOWED_ORIGINS for CORS'),
    ('10th', '#10', 'security.py',       'Change 400 → 403 FORBIDDEN on admin check'),
    ('11th', '#11', 'AdminPage.js',      'Add overflowY:auto + height:100% to Container'),
    ('12th', '#12', 'endpoints.py',      'Move service instantiation into lifespan()'),
    ('13th', '#13', 'chroma_service.py', 'Fix similarity formula (exp(-d) → proper conversion)'),
    ('14th', '#14', 'conftest.py',       'Add service mocks to prevent model loads in tests'),
    ('15th', '#15–22', 'Various',        'Minor cleanup: deprecated imports, pad_token, URLs, dead code'),
]

alt = False
for row_data in ORDER_ROWS:
    row = order_tbl.add_row()
    bg = ALT_BG if alt else 'FFFFFF'
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        shade_cell(row.cells[i], bg)
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        if i == 0:
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    alt = not alt

# ── Save ───────────────────────────────────────────────────────
os.makedirs('output', exist_ok=True)
out_path = 'output/CONFID_AI_Bug_Report.docx'
doc.save(out_path)
print(f"Saved: {out_path}")